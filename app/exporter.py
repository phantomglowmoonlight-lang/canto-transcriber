"""
匯出模組：TXT 與 DOCX 格式輸出
"""
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.speaker_registry import SpeakerRegistry
from app.text_processor import _format_timestamp

logger = logging.getLogger(__name__)


def export_txt(
    segments: list[dict],
    output_path: Path,
    audio_filename: str = "",
    written: bool = False,
    registry: Optional[SpeakerRegistry] = None,
) -> Path:
    """
    匯出純文字 TXT 檔案

    written: True 時匯出書面語版本
    """
    audio_duration = segments[-1]["end"] if segments else 0

    lines = []
    lines.append(f"會議記錄：{audio_filename}")
    lines.append(f"轉寫時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"時長：{_format_timestamp(audio_duration)}")
    lang_label = "書面語" if written else "粵語原文"
    lines.append(f"語言版本：{lang_label}")
    lines.append("")
    lines.append("=" * 40)
    lines.append("")

    for seg in segments:
        timestamp = f"[{_format_timestamp(seg['start'])}]"
        speaker = seg.get("speaker", "人物 #1")
        if registry:
            speaker = registry.get_name(speaker)
        text = seg.get("text_written" if written else "text", "")
        lines.append(f"{timestamp} {speaker}：")
        lines.append(text)
        lines.append("")

    content = "\n".join(lines)
    output_path.write_text(content, encoding="utf-8")
    logger.info(f"TXT 匯出完成：{output_path}")
    return output_path


def export_docx(
    segments: list[dict],
    output_path: Path,
    audio_filename: str = "",
    written: bool = False,
    registry: Optional[SpeakerRegistry] = None,
    font_name: str = "微軟正黑體",
    font_size: int = 12,
    include_timestamp: bool = True,
    include_speaker_color: bool = True,
) -> Path:
    """
    匯出帶格式 DOCX 檔案

    - 說話人姓名：粗體 + 深藍色
    - 時間戳：灰色小字
    - 段落間距：1.15 倍行高
    """
    doc = Document()

    # 設定預設字型
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    # CJK 字型需額外設定 eastAsia 屬性，否則中文回退到預設字型
    try:
        from docx.oxml.ns import qn
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass  # 向後相容，舊版 python-docx 可能不支援
    style.paragraph_format.line_spacing = 1.15

    # 標題
    title = doc.add_heading("會議記錄", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 資訊區
    audio_duration = segments[-1]["end"] if segments else 0
    meta_lines = [
        f"檔案：{audio_filename}",
        f"轉寫時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"時長：{_format_timestamp(audio_duration)}",
        f"語言版本：{'書面語' if written else '粵語原文'}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("")  # 空行

    # 內容區
    SPEAKER_COLORS = [
        RGBColor(0x4F, 0x46, 0xE5),  # 靛藍
        RGBColor(0x7C, 0x3A, 0xED),  # 紫
        RGBColor(0xDB, 0x27, 0x77),  # 粉紅
        RGBColor(0xDC, 0x26, 0x26),  # 紅
        RGBColor(0xEA, 0x58, 0x0C),  # 橙
        RGBColor(0xD9, 0x77, 0x06),  # 琥珀
        RGBColor(0x65, 0xA3, 0x0D),  # 青檸
        RGBColor(0x05, 0x96, 0x69),  # 翠綠
        RGBColor(0x08, 0x91, 0xB2),  # 青
        RGBColor(0x25, 0x63, 0xEB),  # 藍
    ]

    speaker_color_map = {}
    color_idx = 0

    for seg in segments:
        speaker = seg.get("speaker", "人物 #1")
        if registry:
            speaker = registry.get_name(speaker)

        # 時間戳與說話人
        header = f"{_format_timestamp(seg['start'])}  " if include_timestamp else ""
        header += speaker

        p = doc.add_paragraph()
        # 時間戳灰色
        if include_timestamp:
            ts_run = p.add_run(f"{_format_timestamp(seg['start'])}  ")
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # 說話人粗體
        speaker_run = p.add_run(speaker)
        speaker_run.bold = True
        if include_speaker_color:
            if speaker not in speaker_color_map:
                speaker_color_map[speaker] = SPEAKER_COLORS[color_idx % len(SPEAKER_COLORS)]
                color_idx += 1
            speaker_run.font.color.rgb = speaker_color_map[speaker]

        # 內文
        text = seg.get("text_written" if written else "text", "")
        content_p = doc.add_paragraph(text)
        content_p.paragraph_format.left_indent = Inches(0.3)

    doc.save(str(output_path))
    logger.info(f"DOCX 匯出完成：{output_path}")
    return output_path
